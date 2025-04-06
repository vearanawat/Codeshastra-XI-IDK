import os
import shutil
from supabase import create_client, Client
from dotenv import load_dotenv
import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma

# Base directory for application
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Load environment variables
load_dotenv()

# Configuration
CHROMA_PATH = os.path.join(BASE_DIR, "chroma")
SUPABASE_URL = os.getenv("NEXT_PUBLIC_SUPABASE_URL")
SUPABASE_KEY = os.getenv("NEXT_PUBLIC_SUPABASE_ANON_KEY")
BUCKET_NAME = "documents"

print(f"Chroma DB path: {CHROMA_PATH}")
print(f"Supabase URL: {SUPABASE_URL}")
print(f"Bucket name: {BUCKET_NAME}")

# Initialize Supabase client
if not SUPABASE_URL or not SUPABASE_KEY:
    print("Warning: Supabase URL or Key not set in environment variables.")
    supabase = None
else:
    try:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
        print("Supabase client initialized successfully")
    except Exception as e:
        print(f"Error initializing Supabase client: {e}")
        supabase = None

# Create Chroma DB directory if it doesn't exist
os.makedirs(CHROMA_PATH, exist_ok=True)

# Initialize embeddings
try:
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    print("Embeddings model initialized successfully")
except Exception as e:
    print(f"Error initializing embeddings model: {e}")
    raise

# Initialize Chroma DB
try:
    if os.path.exists(CHROMA_PATH) and any(os.listdir(CHROMA_PATH)):
        db = Chroma(persist_directory=CHROMA_PATH, embedding_function=embeddings)
        print("Loaded existing Chroma DB")
    else:
        print("Creating new Chroma DB with initialization document")
        # Create a dummy document to initialize the DB if it doesn't exist
        db = Chroma.from_documents(
            [Document(page_content="Initialization document for Penguin AI.", metadata={"source": "init"})],
            embeddings,
            persist_directory=CHROMA_PATH
        )
        db.persist()
        print("New Chroma DB created and persisted")
except Exception as e:
    print(f"Error initializing Chroma DB: {e}")
    raise


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extracts text from PDF file content."""
    try:
        import io
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""

def extract_text_from_docx(file_content: bytes) -> str:
    """Extracts text from DOCX file content."""
    try:
        import io
        document = DocxDocument(io.BytesIO(file_content))
        text = "\n".join([para.text for para in document.paragraphs])
        return text
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return ""

def extract_text_from_txt(file_content: bytes) -> str:
    """Extracts text from plain text file content."""
    try:
        return file_content.decode('utf-8')
    except Exception as e:
        print(f"Error extracting text from txt file: {e}")
        return ""

def extract_text(file_content: bytes, file_name: str) -> str:
    """Extracts text based on file extension."""
    file_name_lower = file_name.lower()
    if file_name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_content)
    elif file_name_lower.endswith(".docx"):
        return extract_text_from_docx(file_content)
    elif file_name_lower.endswith(".txt"):
        return extract_text_from_txt(file_content)
    else:
        print(f"Unsupported file type: {file_name}")
        return ""

def get_all_files_from_bucket() -> list:
    """Lists all files in the Supabase storage bucket."""
    if not supabase:
        print("Supabase client not initialized. Cannot list files.")
        return []
        
    try:
        all_files = []
        folders = supabase.storage.from_(BUCKET_NAME).list()
        for folder in folders:
            if folder['id'] is None: # It's a directory (user ID)
                user_id = folder['name']
                categories = supabase.storage.from_(BUCKET_NAME).list(user_id)
                for category_folder in categories:
                     if category_folder['id'] is None: # It's a category directory
                        category = category_folder['name']
                        path = f"{user_id}/{category}"
                        files = supabase.storage.from_(BUCKET_NAME).list(path)
                        for file_info in files:
                             if file_info['id'] is not None: # It's a file
                                file_info['user_id'] = user_id
                                file_info['category'] = category
                                file_info['path'] = f"{path}/{file_info['name']}"
                                all_files.append(file_info)
        return all_files
    except Exception as e:
        print(f"Error listing files from bucket: {e}")
        return []

def process_and_store_documents():
    """Downloads files, extracts text, creates chunks, and stores in Chroma."""
    if not supabase:
        print("Supabase client not initialized. Cannot process documents.")
        print("Please set NEXT_PUBLIC_SUPABASE_URL and NEXT_PUBLIC_SUPABASE_ANON_KEY environment variables.")
        return
        
    print("Starting document processing...")
    files_to_process = get_all_files_from_bucket()
    print(f"Found {len(files_to_process)} files in Supabase bucket.")
    
    if not files_to_process:
        print("No files found in Supabase bucket. Please upload some documents first.")
        return
    
    new_chunks_count = 0
    processed_files = set()
    
    # Load existing document sources from Chroma to avoid duplicates
    try:
        existing_docs = db.get(include=["metadatas"])
        processed_files = {metadata['source'] for metadata in existing_docs.get('metadatas', []) if metadata and 'source' in metadata}
        print(f"Found {len(processed_files)} already processed files in Chroma DB.")
    except Exception as e:
        print(f"Warning: Could not retrieve existing documents from Chroma DB: {e}")

    
    for file_info in files_to_process:
        file_path = file_info['path']
        file_name = file_info['name']
        user_id = file_info['user_id']
        category = file_info['category']

        # Skip if already processed
        if file_path in processed_files:
             print(f"Skipping already processed file: {file_path}")
             continue

        print(f"Processing file: {file_path}")
        try:
            # Download file content from Supabase
            response = supabase.storage.from_(BUCKET_NAME).download(file_path)
            file_content = response

            # Extract text
            text = extract_text(file_content, file_name)
            if not text:
                print(f"Could not extract text from {file_name}. Skipping.")
                continue

            # Split text into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=100,
                length_function=len,
                add_start_index=True,
            )
            metadata = {
                "source": file_path,
                "user_id": user_id,
                "category": category,
                "original_filename": file_name 
            }
            chunks = text_splitter.create_documents([text], metadatas=[metadata])
            
            # Add chunks to Chroma DB
            if chunks:
                db.add_documents(chunks)
                new_chunks_count += len(chunks)
                processed_files.add(file_path) # Mark as processed
                print(f"Added {len(chunks)} chunks for {file_name}")

        except Exception as e:
            print(f"Error processing file {file_path}: {e}")
            
    # Persist changes to Chroma DB
    if new_chunks_count > 0:
        print(f"Persisting {new_chunks_count} new chunks to Chroma DB...")
        db.persist()
        print("Chroma DB updated.")
    else:
        print("No new documents to add to Chroma DB.")

if __name__ == "__main__":
    process_and_store_documents() 